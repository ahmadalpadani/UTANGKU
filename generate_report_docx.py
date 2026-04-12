"""
UtangKU - Progress Report DOCX Generator
Kelompok: Aplikasi Bergerak
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
from copy import deepcopy


# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set table cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_section_heading(doc, text, level=1):
    """Add a styled section heading with orange background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Orange background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FF6B00')
    pPr.append(shd)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 107, 0)
    return p


def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    return p


def add_check_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('[v]  ' + text)
    run.font.size = Pt(10)
    return p


def add_bullet_item(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def add_progress_bar(doc, label, pct):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run_label = p.add_run(f'{label} ')
    run_label.font.size = Pt(10)
    run_pct = p.add_run(f'[{pct}%]')
    run_pct.bold = True
    run_pct.font.size = Pt(10)
    if pct == 100:
        run_pct.font.color.rgb = RGBColor(56, 142, 60)
    elif pct == 0:
        run_pct.font.color.rgb = RGBColor(158, 158, 158)
    else:
        run_pct.font.color.rgb = RGBColor(255, 107, 0)
    return p


def add_colored_box(doc, text, bg='FFF3E0', color='FF6B00'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg)
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
    p.paragraph_format.space_after = Pt(4)
    return p


# ─── Main ────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Cover / Title ───────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(10)
title_p.paragraph_format.space_after = Pt(4)
run = title_p.add_run('LAPORAN PROGRESS TUGAS KELOMPOK')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(255, 107, 0)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
run2 = sub_p.add_run('Aplikasi Mobile: UtangKU')
run2.bold = True
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 0, 0)

# Info table
info_table = doc.add_table(rows=7, cols=2)
info_table.style = 'Table Grid'
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Mata Kuliah', 'Aplikasi Bergerak'),
    ('Nama Aplikasi', 'UtangKU'),
    ('Platform', 'Flutter (Dart)'),
    ('Database', 'SQLite (sqflite)'),
    ('State Management', 'Provider'),
    ('Tanggal Laporan', date.today().strftime('%d %B %Y')),
    ('Overall Progress', '~40% Complete'),
]
for i, (key, val) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = val
    set_cell_bg(row.cells[0], 'FFF3E0')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True

# Set column widths
for row in info_table.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(10)

doc.add_paragraph()

# ─── Section 1: Deskripsi Proyek ───────────────────────────────────────────
add_section_heading(doc, '1. Deskripsi Proyek')

add_body(doc,
    'UtangKU adalah aplikasi mobile berbasis Flutter yang dirancang untuk membantu pengguna '
    'mencatat dan mengelola utang dan piutang mereka. Aplikasi ini terintegrasi langsung dengan '
    'WhatsApp untuk memudahkan proses penagihan. Dibangun menggunakan Dart/Flutter dengan SQLite '
    'sebagai database lokal dan Provider untuk state management.')

add_sub_heading(doc, 'Fitur Utama:')
features = [
    'Add/Edit/Delete data utang dan piutang',
    'Pemisahan Utang (yang Anda hutang) dan Piutang (yang orang lain hutang ke Anda)',
    'Filter berdasarkan status pembayaran (Lunas / Belum Lunas)',
    'Deteksi jatuh tempo otomatis dengan peringatan',
    'Integrasi WhatsApp untuk pengiriman pesan tagihan otomatis',
    'Dashboard ringkasan keuangan',
    'Local database storage dengan SQLite',
]
for f in features:
    add_bullet_item(doc, f)

doc.add_paragraph()

# ─── Section 2: Tampilan Aplikasi ──────────────────────────────────────────
add_section_heading(doc, '2. Tampilan Aplikasi')

screens = [
    ('Splash Screen', 'Layar pembuka aplikasi dengan animasi/logo UtangKU'),
    ('Dashboard / Home', 'Halaman utama menampilkan ringkasan saldo, total utang, total piutang, dan transaksi terbaru'),
    ('Daftar Utang', 'Halaman daftar utang dengan filter status dan aksi cepat (tagih via WA)'),
    ('Daftar Piutang', 'Halaman daftar piutang dengan filter dan integrasi WhatsApp'),
    ('Form Tambah / Edit', 'Form input untuk menambah atau mengedit data utang/piutang'),
    ('Settings', 'Pengaturan aplikasi'),
]

screen_table = doc.add_table(rows=len(screens), cols=2)
screen_table.style = 'Table Grid'
for i, (name, desc) in enumerate(screens):
    row = screen_table.rows[i]
    row.cells[0].text = name
    row.cells[1].text = desc
    if i % 2 == 0:
        set_cell_bg(row.cells[0], 'FFF3E0')
        set_cell_bg(row.cells[1], 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].width = Cm(4.5)

doc.add_paragraph()

# ─── Section 3: Fitur yang Sudah Dibangun ───────────────────────────────────
add_section_heading(doc, '3. Fitur yang Sudah Dibangun (Phase 1 & 2)')

# Phase 1
add_sub_heading(doc, '[SELESAI] Phase 1: Setup & Foundation (100%)')
for item in [
    'Struktur proyek Flutter dengan folder yang terorganisir',
    'Konfigurasi pubspec.yaml dengan semua dependencies',
    'Setup tema Material 3 dengan warna orange (#FF6B00)',
    'Database helper SQLite dengan schema lengkap',
    'Model data (DebtModel, DebtType, PaymentStatus)',
    'Provider state management',
    'WhatsApp service untuk generate pesan tagihan',
    'Formatters untuk mata uang IDR dan tanggal lokal Indonesia',
]:
    add_check_item(doc, item)

doc.add_paragraph()

# Phase 2
add_sub_heading(doc, '[SELESAI] Phase 2: Full CRUD Functionality (100%)')
for item in [
    'Add Debt Screen - Form lengkap dengan validasi, date picker, live preview',
    'Debt List Screen - Daftar utang dengan filter (Semua/Belum Lunas/Lunas)',
    'Piutang List Screen - Daftar piutang dengan filter dan aksi sama',
    'WhatsApp Integration - Tombol "Tagih via WhatsApp" dengan pesan otomatis',
    'Toggle Status - Tandai Lunas / Belum Lunas dengan konfirmasi',
    'Long-press menu - Edit dan Delete dengan dialog konfirmasi',
    'Overdue detection - Peringatan otomatis untuk yang lewat jatuh tempo',
    'Empty state handling - Tampilan khusus ketika belum ada data',
]:
    add_check_item(doc, item)

doc.add_paragraph()

# ─── Section 4: Struktur Kode ───────────────────────────────────────────────
add_section_heading(doc, '4. Struktur Kode / File')

add_body(doc, 'Direktori lib/', bold=True)

code_table = doc.add_table(rows=len(structure_data := [
    ('lib/main.dart', 'Entry point aplikasi'),
    ('lib/models/', 'Data models (debt_model.dart, debt_type.dart, payment_status.dart)'),
    ('lib/database/', 'Database helper & service (SQLite CRUD)'),
    ('lib/screens/home/', 'Splash screen & Dashboard/Home screen'),
    ('lib/screens/debt/', 'Add debt screen & Debt list screen'),
    ('lib/screens/piutang/', 'Piutang list screen'),
    ('lib/screens/settings/', 'Settings screen'),
    ('lib/services/', 'WhatsApp service & Mock service'),
    ('lib/utils/', 'Theme, constants, formatters (IDR currency, tanggal Indonesia)'),
    ('lib/providers/', 'DebtProvider - State management'),
]), cols=2)
code_table.style = 'Table Grid'
for i, (path, desc) in enumerate(structure_data):
    row = code_table.rows[i]
    row.cells[0].text = path
    row.cells[1].text = desc
    if i % 2 == 0:
        set_cell_bg(row.cells[0], 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = 'Courier New'
    row.cells[0].width = Cm(5)

doc.add_paragraph()

# Lines of code
add_sub_heading(doc, 'Total Lines of Code (Phase 1 & 2):')
for line in [
    'add_debt_screen.dart      : ~380 baris',
    'debt_list_screen.dart     : ~430 baris',
    'piutang_list_screen.dart  : ~440 baris',
    'home_screen.dart + splash : ~200 baris',
    'database_helper.dart      : ~150 baris',
    'Total                     : ~1.600+ baris kode production-ready!',
]:
    add_body(doc, '  - ' + line)

doc.add_paragraph()

# ─── Section 5: Technology Stack ──────────────────────────────────────────────
add_section_heading(doc, '5. Technology Stack')

stack_table = doc.add_table(rows=1 + len(stack_data := [
    ('Framework', 'Flutter (Dart)', 'Cross-platform mobile development'),
    ('Language', 'Dart', 'Bahasa pemrograman utama'),
    ('Database', 'SQLite (sqflite ^2.3.0)', 'Local database storage'),
    ('State Mgmt', 'Provider (^6.1.1)', 'State management'),
    ('UI', 'Material Design 3', 'Desain antarmuka pengguna'),
    ('Charts', 'fl_chart (^0.66.0)', 'Visualisasi data (prepared)'),
    ('WhatsApp', 'url_launcher + whatsapp_unilink', 'Integrasi WhatsApp'),
    ('Date/Time', 'intl + table_calendar', 'Formatting & kalender'),
    ('Icons', 'cupertino_icons', 'Ikon iOS-style'),
]), cols=3)
stack_table.style = 'Table Grid'

# Header row
hdr = stack_table.rows[0]
hdr.cells[0].text = 'Kategori'
hdr.cells[1].text = 'Teknologi'
hdr.cells[2].text = 'Keterangan'
set_cell_bg(hdr.cells[0], 'FF6B00')
set_cell_bg(hdr.cells[1], 'FF6B00')
set_cell_bg(hdr.cells[2], 'FF6B00')
for cell in hdr.cells:
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

for i, (cat, tech, ket) in enumerate(stack_data):
    row = stack_table.rows[i + 1]
    row.cells[0].text = cat
    row.cells[1].text = tech
    row.cells[2].text = ket
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(9)

doc.add_paragraph()

# ─── Section 6: Progress Chart ───────────────────────────────────────────────
add_section_heading(doc, '6. Grafik Progress per Phase')

doc.add_paragraph()
for label, pct in [
    ('Phase 1: Setup & Foundation', 100),
    ('Phase 2: Full CRUD', 100),
    ('Phase 3: Statistics & Charts', 0),
    ('Phase 4: Notifications', 0),
    ('Phase 5: Security (PIN/Biometric)', 0),
]:
    add_progress_bar(doc, label, pct)
doc.add_paragraph()

# Overall
overall_p = doc.add_paragraph()
overall_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
overall_run = overall_p.add_run('Overall Progress: ~40% Complete')
overall_run.bold = True
overall_run.font.size = Pt(12)
overall_run.font.color.rgb = RGBColor(255, 107, 0)

doc.add_paragraph()

# ─── Section 7: Next Steps ───────────────────────────────────────────────────
add_section_heading(doc, '7. Rencana Selanjutnya (Next Steps)')

next_steps = [
    ('Statistics & Charts (Phase 3) - [0%]', [
        'Implementasi fl_chart untuk visualisasi data',
        'Monthly summary chart',
        'Category breakdown chart',
        'Visual trend analysis',
    ]),
    ('Notifications (Phase 4) - [0%]', [
        'Due date reminders / notifikasi',
        'Notification permissions handling',
        'Scheduled notifications',
    ]),
    ('Security / PIN (Phase 5) - [0%]', [
        'PIN setup screen',
        'Biometric authentication (Face ID / Fingerprint)',
        'Auto-lock after timeout',
    ]),
    ('Polish & Extra Features - [0%]', [
        'Search functionality',
        'Export data (CSV/PDF)',
        'Dark mode',
        'Settings screen improvements',
    ]),
]

for title, items in next_steps:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3E0')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(56, 142, 60)
    for item in items:
        add_bullet_item(doc, item)
    doc.add_paragraph()

doc.add_paragraph()

# ─── Section 8: User Flow ───────────────────────────────────────────────────
add_section_heading(doc, '8. Alur Penggunaan (User Flow)')

user_flows = [
    ('Menambah Utang/Piutang',
     'Dashboard > Tap "Tambah Utang/Piutang" > Isi form > Simpan > Data tersimpan'),
    ('Melihat Daftar',
     'Bottom Nav > Tab "Utang/Piutang" > Lihat daftar > Filter sesuai kebutuhan'),
    ('Menandai Lunas',
     'Daftar Utang/Piutang > Tap "Tandai Lunas" > Konfirmasi > Status berubah hijau'),
    ('Edit / Hapus Data',
     'Tekan lama (long-press) pada item > Pilih "Edit" atau "Hapus" > Konfirmasi'),
    ('Tagih via WhatsApp',
     'Daftar Piutang > Tap tombol "Tagih" > WhatsApp terbuka dengan pesan otomatis'),
    ('Filter Daftar',
     'Tap chip filter: "Semua" | "Belum Lunas" | "Lunas" > Daftar terfilter'),
]

flow_table = doc.add_table(rows=len(user_flows), cols=2)
flow_table.style = 'Table Grid'
for i, (aksi, flow) in enumerate(user_flows):
    row = flow_table.rows[i]
    row.cells[0].text = aksi
    row.cells[1].text = flow
    if i % 2 == 0:
        set_cell_bg(row.cells[0], 'FFF3E0')
        set_cell_bg(row.cells[1], 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].width = Cm(4.5)

doc.add_paragraph()

# ─── Section 9: Warna & Desain ──────────────────────────────────────────────
add_section_heading(doc, '9. Warna & Desain UI')

color_table = doc.add_table(rows=1 + len(color_data := [
    ('Primary / Accent', '#FF6B00', 'Orange', 'Warna utama aplikasi, tombol, aksen'),
    ('Utang Color', '#E53935', 'Red', 'Indikator utang (uang yang Anda hutang)'),
    ('Piutang Color', '#43A047', 'Green', 'Indikator piutang (uang yang orang hutang ke Anda)'),
    ('Success', '#388E3C', 'Green', 'Status lunas, konfirmasi sukses'),
    ('Warning', '#F57C00', 'Orange', 'Peringatan jatuh tempo'),
    ('Error', '#D32F2F', 'Red', 'Error, hapus data'),
    ('Background', '#FAFAFA', 'Light Gray', 'Background halaman'),
    ('Card BG', '#FFFFFF', 'White', 'Kartu komponen'),
]), cols=4)
color_table.style = 'Table Grid'

hdr = color_table.rows[0]
hdr.cells[0].text = 'Nama Warna'
hdr.cells[1].text = 'Hex Code'
hdr.cells[2].text = 'Warna'
hdr.cells[3].text = 'Keterangan'
for cell in hdr.cells:
    set_cell_bg(cell, 'FF6B00')
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

for i, (name, hex_code, color_name, desc) in enumerate(color_data):
    row = color_table.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = hex_code
    row.cells[2].text = color_name
    row.cells[3].text = desc
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(9)

doc.add_paragraph()

# ─── Section 10: Kesimpulan ───────────────────────────────────────────────────
add_section_heading(doc, '10. Kesimpulan')

add_body(doc,
    'Aplikasi UtangKU telah berhasil dibangun dengan fitur utama CRUD (Create, Read, Update, Delete) '
    'yang berjalan dengan baik. Fase 1 (Setup) dan Fase 2 (CRUD) telah selesai 100%. '
    'Aplikasi ini sudah dapat digunakan secara fungsional untuk mencatat dan mengelola utang serta piutang.')
add_body(doc,
    'Fitur integrasi WhatsApp memungkinkan pengguna untuk langsung mengirim pesan tagihan hanya dengan '
    'satu ketukan. Dashboard memberikan ringkasan cepat kondisi keuangan. Aplikasi siap untuk '
    'dikembangkan lebih lanjut dengan fitur statistik, notifikasi, dan keamanan (PIN/Biometric).')

doc.add_paragraph()

# Ringkasan box
add_colored_box(doc, 'RINGKASAN PROGRESS', bg='FFF3E0', color='FF6B00')
ringkasan_headers = ['Phase', 'Progress', 'Status']
ringkasan_data = [
    ('Phase 1 (Setup)', '100%', 'Selesai'),
    ('Phase 2 (CRUD)', '100%', 'Selesai'),
    ('Phase 3 (Stats)', '0%', 'Belum dimulai'),
    ('Phase 4 (Notif)', '0%', 'Belum dimulai'),
    ('Phase 5 (Security)', '0%', 'Belum dimulai'),
]
ringkasan_table = doc.add_table(rows=1 + len(ringkasan_data), cols=3)
ringkasan_table.style = 'Table Grid'

hdr_row = ringkasan_table.rows[0]
for i, h in enumerate(ringkasan_headers):
    hdr_row.cells[i].text = h
    set_cell_bg(hdr_row.cells[i], 'FF6B00')
    for para in hdr_row.cells[i].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

for i, (name, pct, status) in enumerate(ringkasan_data):
    row = ringkasan_table.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = pct
    row.cells[2].text = status
    if i % 2 == 0:
        for cell in row.cells:
            set_cell_bg(cell, 'F5F5F5')
    for cell in row.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
    # Color pct cell
    if pct == '100%':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(56, 142, 60)
    else:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(158, 158, 158)
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(2.5)
    row.cells[2].width = Cm(4)

doc.add_paragraph()

# ─── Output ─────────────────────────────────────────────────────────────────
output_path = '/Users/nadya/Documents/SEMESTER_4/aplikasi_bergerak/utangku_app/UtangKU_Progress_Report.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
